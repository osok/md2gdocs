#!/usr/bin/env python3
"""
Convert Markdown with Mermaid diagrams to Microsoft Word (DOCX).

This script parses markdown files, renders mermaid diagrams as images,
and creates a DOCX file with the content and embedded diagram images.

Requirements:
    pip install python-docx
    pip install markdown
    pip install Pillow
    pip install requests

    Also requires either:
    - A local mermaid CLI installation (npm install -g @mermaid-js/mermaid-cli)
    - Or uses the mermaid.ink API for rendering
"""

import os
import re
import base64
import tempfile
import subprocess
import argparse
from pathlib import Path
from typing import List, Dict, Tuple, Optional
import time

# DOCX imports
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Image handling
from PIL import Image
import requests


class MermaidRenderer:
    """Handle rendering of Mermaid diagrams to images."""

    def __init__(self, use_api: bool = True):
        """
        Initialize the renderer.

        Args:
            use_api: If True, use mermaid.ink API. If False, use local mermaid CLI.
        """
        self.use_api = use_api
        if not use_api:
            # Check if mermaid CLI is installed
            try:
                subprocess.run(['mmdc', '--version'], check=True, capture_output=True)
            except (subprocess.CalledProcessError, FileNotFoundError):
                print("Warning: mermaid CLI not found. Falling back to API.")
                self.use_api = True

    def render_to_image(self, mermaid_code: str, output_path: str) -> bool:
        """
        Render mermaid code to an image file.

        Args:
            mermaid_code: The mermaid diagram code
            output_path: Path where the image should be saved

        Returns:
            True if successful, False otherwise
        """
        if self.use_api:
            return self._render_with_api(mermaid_code, output_path)
        else:
            return self._render_with_cli(mermaid_code, output_path)

    def _render_with_api(self, mermaid_code: str, output_path: str) -> bool:
        """Render using mermaid.ink API with retry logic."""
        max_retries = 3
        retry_delay = 2  # seconds

        for attempt in range(max_retries):
            try:
                # Encode the mermaid code for the API
                encoded = base64.urlsafe_b64encode(
                    mermaid_code.encode('utf-8')
                ).decode('ascii')

                # Request the image from mermaid.ink
                url = f"https://mermaid.ink/img/{encoded}"
                response = requests.get(url, timeout=30)
                response.raise_for_status()

                # Save the image
                with open(output_path, 'wb') as f:
                    f.write(response.content)

                return True
            except requests.exceptions.HTTPError as e:
                if attempt < max_retries - 1 and e.response.status_code in [503, 429, 500]:
                    # Retry on server errors or rate limiting
                    print(f"Mermaid API error (attempt {attempt + 1}/{max_retries}): {e.response.status_code}. Retrying in {retry_delay}s...")
                    time.sleep(retry_delay)
                    retry_delay *= 2  # Exponential backoff
                else:
                    print(f"Error rendering mermaid diagram with API: {e}")
                    return False
            except Exception as e:
                print(f"Error rendering mermaid diagram with API: {e}")
                return False

        return False

    def _render_with_cli(self, mermaid_code: str, output_path: str) -> bool:
        """Render using local mermaid CLI."""
        try:
            # Create temporary file with mermaid code
            with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as f:
                f.write(mermaid_code)
                temp_mmd = f.name

            # Run mermaid CLI
            subprocess.run([
                'mmdc',
                '-i', temp_mmd,
                '-o', output_path,
                '-t', 'default',
                '-b', 'white'
            ], check=True, capture_output=True)

            # Clean up
            os.unlink(temp_mmd)
            return True
        except Exception as e:
            print(f"Error rendering mermaid diagram with CLI: {e}")
            if 'temp_mmd' in locals() and os.path.exists(temp_mmd):
                os.unlink(temp_mmd)
            return False


class MarkdownToDocx:
    """Convert Markdown with Mermaid diagrams to DOCX."""

    def __init__(self):
        """Initialize the converter."""
        self.mermaid_renderer = MermaidRenderer(use_api=True)

    def parse_markdown(self, markdown_text: str) -> Tuple[List[Dict], List[str]]:
        """
        Parse markdown and extract mermaid diagrams and tables.

        Args:
            markdown_text: The markdown content

        Returns:
            Tuple of (content blocks, mermaid codes)
        """
        blocks = []
        mermaid_diagrams = []

        # First, extract tables before splitting by code blocks
        # Tables are identified by pipes - match all consecutive lines with pipes
        table_pattern = r'(\|.+\|(?:\n\|.+\|)+)'

        # Replace tables with placeholders
        table_matches = []
        def replace_table(match):
            table_matches.append(match.group(0))
            return f'<<<TABLE_{len(table_matches) - 1}>>>'

        markdown_with_placeholders = re.sub(table_pattern, replace_table, markdown_text)

        # Split by code blocks
        parts = re.split(r'```(\w+)?\n(.*?)```', markdown_with_placeholders, flags=re.DOTALL)

        for i, part in enumerate(parts):
            if i % 3 == 0:  # Regular markdown content
                if part.strip():
                    # Check for table placeholders
                    table_placeholder_pattern = r'<<<TABLE_(\d+)>>>'
                    sub_parts = re.split(table_placeholder_pattern, part)

                    for j, sub_part in enumerate(sub_parts):
                        if j % 2 == 0:  # Regular markdown
                            if sub_part.strip():
                                blocks.append({'type': 'markdown', 'content': sub_part})
                        else:  # Table index
                            table_idx = int(sub_part)
                            blocks.append({'type': 'table', 'content': table_matches[table_idx]})
            elif i % 3 == 1:  # Code block language
                lang = part
            else:  # Code block content
                if lang and lang.lower() == 'mermaid':
                    mermaid_diagrams.append(part)
                    blocks.append({'type': 'mermaid', 'index': len(mermaid_diagrams) - 1})
                else:
                    # Other code blocks
                    blocks.append({'type': 'code', 'language': lang or '', 'content': part})

        return blocks, mermaid_diagrams

    def create_docx(self, title: str, blocks: List[Dict], mermaid_images: List[str]) -> Document:
        """
        Create a DOCX document with the parsed content.

        Args:
            title: Document title
            blocks: Content blocks from parsing
            mermaid_images: Paths to rendered mermaid images

        Returns:
            The Document object
        """
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        for block in blocks:
            if block['type'] == 'markdown':
                self._add_markdown_to_doc(doc, block['content'])
            elif block['type'] == 'code':
                self._add_code_block(doc, block['content'], block['language'])
            elif block['type'] == 'table':
                self._add_table_to_doc(doc, block['content'])
            elif block['type'] == 'mermaid':
                image_path = mermaid_images[block['index']]
                if os.path.exists(image_path):
                    doc.add_picture(image_path, width=Inches(6))
                    doc.add_paragraph()  # Add spacing after image

        return doc

    def _add_markdown_to_doc(self, doc: Document, markdown_text: str):
        """
        Add markdown content to the document with proper formatting.

        Args:
            doc: The Document object
            markdown_text: Markdown content to add
        """
        lines = markdown_text.split('\n')

        for line in lines:
            # Handle headers
            header_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2)
                para = doc.add_heading(level=level)
                self._add_formatted_text(para, text)
                continue

            # Handle list items
            list_match = re.match(r'^(\s*)[\*\-]\s+(.+)$', line)
            if list_match:
                indent = list_match.group(1)
                text = list_match.group(2)
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_text(para, text)
                # Set indentation level
                para.paragraph_format.left_indent = Inches(len(indent) * 0.25)
                continue

            # Handle numbered lists
            num_list_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
            if num_list_match:
                indent = num_list_match.group(1)
                text = num_list_match.group(3)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_text(para, text)
                para.paragraph_format.left_indent = Inches(len(indent) * 0.25)
                continue

            # Regular paragraph
            if line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line)
            else:
                # Empty line
                doc.add_paragraph()

    def _add_formatted_text(self, paragraph, text: str):
        """
        Add text to a paragraph with inline formatting (bold, italic).

        Args:
            paragraph: The paragraph object
            text: Text with markdown formatting
        """
        # Find all formatting markers
        # This regex finds bold (**text**), italic (*text*), and regular text
        pattern = r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_|\[.*?\]\(.*?\)|[^*_\[]+)'

        parts = re.findall(pattern, text)

        for part in parts:
            if not part:
                continue

            # Bold: **text** or __text__
            if (part.startswith('**') and part.endswith('**')) or \
               (part.startswith('__') and part.endswith('__')):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            # Italic: *text* or _text_ (but not part of bold)
            elif (part.startswith('*') and part.endswith('*') and not part.startswith('**')) or \
                 (part.startswith('_') and part.endswith('_') and not part.startswith('__')):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            # Links: [text](url)
            elif part.startswith('[') and '](' in part:
                link_match = re.match(r'\[(.+?)\]\((.+?)\)', part)
                if link_match:
                    run = paragraph.add_run(link_match.group(1))
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
            else:
                # Regular text
                paragraph.add_run(part)

    def _add_code_block(self, doc: Document, code: str, language: str):
        """
        Add a code block with professional formatting.

        Args:
            doc: The Document object
            code: The code content
            language: Programming language
        """
        # Add a paragraph with code styling
        para = doc.add_paragraph()
        run = para.add_run(code)

        # Format as code
        run.font.name = 'Courier New'
        run.font.size = Pt(10)

        # Set paragraph background (shading)
        from docx.oxml.shared import OxmlElement
        from docx.oxml.ns import qn

        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F2F2F2')  # Light grey background
        para._p.get_or_add_pPr().append(shd)

        # Add border
        from docx.oxml import parse_xml
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')  # Border size
            border.set(qn('w:space'), '10')  # Padding
            border.set(qn('w:color'), '000000')  # Black
            pBdr.append(border)

        pPr.append(pBdr)

    def _add_table_to_doc(self, doc: Document, table_text: str):
        """
        Add a markdown table to the document.

        Args:
            doc: The Document object
            table_text: Markdown table text
        """
        # Parse table rows
        lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]

        if len(lines) < 2:
            return  # Invalid table

        # Parse header
        header_cells = [cell.strip() for cell in lines[0].split('|') if cell.strip()]

        # Skip separator line (line 1)

        # Parse data rows
        data_rows = []
        for line in lines[2:]:
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            if cells:
                data_rows.append(cells)

        # Create table
        table = doc.add_table(rows=1 + len(data_rows), cols=len(header_cells))
        table.style = 'Light Grid Accent 1'

        # Add header row
        header_row = table.rows[0]
        for i, header_text in enumerate(header_cells):
            cell = header_row.cells[i]
            # Remove markdown formatting from header
            clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', header_text)
            clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
            cell.text = clean_text

            # Format header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)

            # Set background color for header
            from docx.oxml.shared import OxmlElement
            from docx.oxml.ns import qn
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'D9E2F3')  # Light blue
            cell._element.get_or_add_tcPr().append(shd)

        # Add data rows
        for row_idx, row_data in enumerate(data_rows):
            row = table.rows[row_idx + 1]
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < len(header_cells):  # Ensure we don't exceed columns
                    cell = row.cells[col_idx]
                    # Remove markdown formatting and handle code blocks
                    clean_text = re.sub(r'`(.+?)`', r'\1', cell_text)
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_text)
                    cell.text = clean_text

                    # Format code-like content (backticks)
                    if '`' in cell_text:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)

        # Add spacing after table
        doc.add_paragraph()

    def convert(self, markdown_file: str, output_file: Optional[str] = None) -> str:
        """
        Convert a markdown file to DOCX.

        Args:
            markdown_file: Path to the markdown file
            output_file: Optional output file path

        Returns:
            The output file path
        """
        # Read markdown file
        with open(markdown_file, 'r', encoding='utf-8') as f:
            markdown_content = f.read()

        # Parse markdown and extract mermaid diagrams
        blocks, mermaid_codes = self.parse_markdown(markdown_content)

        # Render mermaid diagrams
        mermaid_images = []
        with tempfile.TemporaryDirectory() as temp_dir:
            for i, mermaid_code in enumerate(mermaid_codes):
                image_path = os.path.join(temp_dir, f'mermaid_{i}.png')
                if self.mermaid_renderer.render_to_image(mermaid_code, image_path):
                    mermaid_images.append(image_path)
                else:
                    mermaid_images.append('')

            # Create DOCX
            title = Path(markdown_file).stem
            doc = self.create_docx(title, blocks, mermaid_images)

            # Determine output file path
            if not output_file:
                # Create docx directory in the same location as the markdown file
                md_path = Path(markdown_file)
                docx_dir = md_path.parent / 'docx'
                docx_dir.mkdir(exist_ok=True)
                output_file = str(docx_dir / f"{md_path.stem}.docx")

            # Ensure parent directory exists
            Path(output_file).parent.mkdir(parents=True, exist_ok=True)

            # Save the document
            doc.save(output_file)

        print(f"\nDocument created successfully: {output_file}")
        return output_file

    def convert_directory(self, directory: str, output_dir: Optional[str] = None) -> List[str]:
        """
        Convert all markdown files in a directory to DOCX.
        Does not process subdirectories.

        Args:
            directory: Path to the directory containing markdown files
            output_dir: Optional output directory (defaults to 'docx' subdirectory)

        Returns:
            List of created DOCX file paths
        """
        # Get all .md files in the directory (not subdirectories)
        directory_path = Path(directory)
        if not directory_path.is_dir():
            raise ValueError(f"'{directory}' is not a valid directory")

        md_files = list(directory_path.glob('*.md'))

        if not md_files:
            print(f"No markdown files found in '{directory}'")
            return []

        print(f"Found {len(md_files)} markdown file(s) in '{directory}'")

        # Set output directory - default to 'docx' subdirectory
        if output_dir:
            output_path = Path(output_dir)
        else:
            output_path = directory_path / 'docx'

        output_path.mkdir(parents=True, exist_ok=True)

        docx_files = []
        for md_file in md_files:
            try:
                print(f"\nProcessing: {md_file.name}")
                output_file = str(output_path / md_file.with_suffix('.docx').name)
                result = self.convert(str(md_file), output_file)
                docx_files.append(result)
            except Exception as e:
                print(f"Error processing {md_file.name}: {e}")
                continue

        print(f"\n{'='*60}")
        print(f"Completed: {len(docx_files)} of {len(md_files)} files converted successfully")
        print(f"Output directory: {output_path}")

        return docx_files


def main():
    """Main entry point."""
    parser = argparse.ArgumentParser(
        description='Convert Markdown with Mermaid diagrams to DOCX'
    )
    parser.add_argument(
        'path',
        help='Path to a markdown file or directory containing markdown files'
    )
    parser.add_argument(
        '-o', '--output',
        help='Output file or directory (default: same location as input with .docx extension)'
    )
    parser.add_argument(
        '--use-cli',
        action='store_true',
        help='Use local mermaid CLI instead of API for rendering'
    )

    args = parser.parse_args()

    # Check if path exists
    if not os.path.exists(args.path):
        print(f"Error: Path '{args.path}' not found")
        return 1

    # Create converter
    converter = MarkdownToDocx()

    # Set rendering method
    if args.use_cli:
        converter.mermaid_renderer = MermaidRenderer(use_api=False)

    try:
        # Check if path is a directory or file
        if os.path.isdir(args.path):
            # Convert all markdown files in the directory
            converter.convert_directory(args.path, args.output)
        else:
            # Convert single file
            converter.convert(args.path, args.output)
        return 0
    except Exception as e:
        print(f"Error: {e}")
        return 1


if __name__ == '__main__':
    exit(main())
